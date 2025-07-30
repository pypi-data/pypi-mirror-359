import os
import time
import zipfile
from io import BytesIO
from docx import Document
from docx.shared import Inches
import openpyxl
import ast
import xml.etree.ElementTree as ET
import re

def convert_to_number_if_possible(value):
    try:
        evaluated_value = ast.literal_eval(value)
        if isinstance(evaluated_value, (int, float)):
            return evaluated_value
    except (ValueError, SyntaxError):
        pass
    return value

def replace_text_in_paragraph(paragraph, replace_text_dict, replace_image_dict, state_dir):
    inline = paragraph.runs
    _stack = ''
    _start_stack = 0
    _choose_item = []
    for item in inline:
        text_inline = item.text
        pattern = r'\{\{(.*?)\}\}'
        matches = re.findall(pattern, text_inline)
        for placeholder in matches:
            # Kiá»ƒm tra náº¿u placeholder cĂ³ trong replace_text_dict
            if placeholder in replace_text_dict:
                # Thay tháº¿ placeholder báº±ng giĂ¡ trá»‹ tÆ°Æ¡ng á»©ng trong dictionary
                replacement_text = replace_text_dict[placeholder]
                text_inline = text_inline.replace("{{"+f"{placeholder}"+"}}", replacement_text)
                item.text = text_inline

        if "{" in text_inline and "}}" not in text_inline.split("{")[-1]:
            _start_stack += text_inline.split("}")[-1].count('{')
        if "}" in text_inline and "{{" not in text_inline.split("}")[0]:
            _choose_item.append(item)
            if _start_stack > 0:
                _start_stack -= text_inline.split("{")[-1].count('}')
                _stack += text_inline.split("}")[0] + "}"*text_inline.split("{")[-1].count('}')
                if _start_stack == 0:
                    placeholder = _stack.split("{{")[1].split("}}")[0]
                    if placeholder in replace_text_dict.keys():
                        replace_item = _choose_item[1]
                        replace_item.text = replace_text_dict[placeholder]
                        _choose_item[-1].text = _choose_item[-1].text.replace("}", "")
                        _choose_item[0].text = _choose_item[0].text.replace("{", "")
                        for i in range(2, len(_choose_item)-2):
                            _choose_item[i].text = ""
                        _stack = ''
                        _choose_item = []

                    if placeholder in replace_image_dict.keys():
                        run = paragraph.add_run()
                        image_path = replace_image_dict[placeholder]
                        image_path = os.path.normpath(image_path).lstrip(os.sep)
                        full_image_path = os.path.join(state_dir, image_path)
                        run.add_picture(full_image_path, width=Inches(6))
                        _choose_item[-1].text = _choose_item[-1].text.replace("}", "")
                        _choose_item[0].text = _choose_item[0].text.replace("{", "")
                        for i in range(1, len(_choose_item)-2):
                            _choose_item[i].text = ""
                        _stack = ''
                        _choose_item = []

        if _start_stack >0 :
            _choose_item.append(item)
            _stack += text_inline

def replace_text_in_excel_sheet(sheet, replace_text_dict):
    for row in sheet.iter_rows():
        for cell in row:
            if cell.value and isinstance(cell.value, str) and '{{' in cell.value and '}}' in cell.value:
                placeholder = cell.value.split('{{')[1].split('}}')[0]
                if placeholder in replace_text_dict:
                    cell.value = convert_to_number_if_possible(replace_text_dict[placeholder])

def extract_and_modify_embedded_excel(docx_zip, embedded_file_name, text_content):
    with docx_zip.open(embedded_file_name) as embedded_file:
        with BytesIO(embedded_file.read()) as excel_file:
            wb = openpyxl.load_workbook(excel_file, data_only=True)
            for sheet in wb.sheetnames:
                worksheet = wb[sheet]
                replace_text_in_excel_sheet(worksheet, text_content)
            modified_excel = BytesIO()
            wb.save(modified_excel)
            return modified_excel.getvalue()

def check_unfill_in_excel_sheet(sheet, replace_text_dict):
    unfill = []
    for row in sheet.iter_rows():
        for cell in row:
            if cell.value and isinstance(cell.value, str) and '{{' in cell.value and '}}' in cell.value:
                placeholder = cell.value.split('{{')[1].split('}}')[0]
                unfill.append(placeholder)
    return unfill
def check_unfill_embedded_excel(docx_zip, embedded_file_name, text_content):
    unfill_list = []
    with docx_zip.open(embedded_file_name) as embedded_file:
        with BytesIO(embedded_file.read()) as excel_file:
            wb = openpyxl.load_workbook(excel_file, data_only=True)
            for sheet in wb.sheetnames:
                worksheet = wb[sheet]
                unfill = check_unfill_in_excel_sheet(worksheet, text_content)
            unfill_list += unfill
            return unfill_list

def check_unfill(paragraph):
    inline = paragraph.runs
    _stack = ''
    _start_stack = 0
    _choose_item = []
    unfill = []
    for item in inline:
        text_inline = item.text
        pattern = r'\{\{(.*?)\}\}'
        matches = re.findall(pattern, text_inline)
        for placeholder in matches:
            unfill.append(placeholder)

        if "{" in text_inline and "}}" not in text_inline.split("{")[-1]:
            _start_stack += text_inline.split("}")[-1].count('{')
        if "}" in text_inline and "{{" not in text_inline.split("}")[0]:
            _choose_item.append(item)
            if _start_stack > 0:
                _start_stack -= text_inline.split("{")[-1].count('}')
                _stack += text_inline.split("}")[0] + "}"*text_inline.split("{")[-1].count('}')
                if _start_stack == 0:
                    placeholder = _stack.split("{{")[1].split("}}")[0]
                    unfill.append(placeholder)

        if _start_stack >0 :
            _choose_item.append(item)
            _stack += text_inline
    return unfill

def extract_and_modify_docx(file_path, output_file_path, text_content, image_content, state_dir):

    unfill_list = []

    with zipfile.ZipFile(file_path, 'r') as docx_zip:
        with zipfile.ZipFile(output_file_path, 'w') as new_docx_zip:
            for item in docx_zip.infolist():
                if item.filename.startswith('word/embeddings/') and item.filename.endswith('.xlsx'):
                    modified_excel_data = extract_and_modify_embedded_excel(docx_zip, item.filename, text_content)
                    unfill_list += check_unfill_embedded_excel(docx_zip, item.filename, text_content)
                    new_docx_zip.writestr(item.filename, modified_excel_data)
                else:
                    new_docx_zip.writestr(item.filename, docx_zip.read(item.filename))

    doc = Document(output_file_path)

    for paragraph in doc.paragraphs:
        replace_text_in_paragraph(paragraph, text_content, image_content, state_dir)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    replace_text_in_paragraph(paragraph, text_content, image_content, state_dir)

    doc.save(output_file_path)

    doc = Document(output_file_path)
    
    for paragraph in doc.paragraphs:
        unfill = check_unfill(paragraph)
        unfill_list += unfill
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    unfill = check_unfill(paragraph)
                    unfill_list += unfill
    
    return True, unfill_list

def execute(file_path: str, output_file_path: str, text_content={}, image_content={}, state_dir: str = None, **kwargs):
    try:
        if len(text_content) == 0 and len(image_content) == 0:
            return {
                "tool": "docx_fill_with_image",
                "success": False,
                "error": "Try again with required \"text_content\":{\"placeholder\":\"value\"} hoáº·c \"image_content\":{\"placeholder\":\"image_path\"}"
            }

        if state_dir is None:
            timestamp = str(int(time.time()))
            state_dir = os.path.join("./test_environment", timestamp)
            os.makedirs(state_dir, exist_ok=True)

        file_path = os.path.normpath(file_path).lstrip(os.sep)
        full_file_path = os.path.join(state_dir, file_path)

        output_file_path = os.path.normpath(output_file_path).lstrip(os.sep)
        full_output_file_path = os.path.join(state_dir, output_file_path)

        _, unfill_list = extract_and_modify_docx(full_file_path, full_output_file_path, text_content, image_content, state_dir)

        for key, img_path in image_content.items():
            text_content[key] = img_path
        return {
            "tool": "docx_fill_with_image",
            "success": True,
            "output_file_path": output_file_path,
            "filled": {**text_content, **image_content},
            "unfilled_placeholder": unfill_list
        }
    except Exception as e:
        return {"tool": "docx_fill_with_image", "success": False, "error": str(e)}
