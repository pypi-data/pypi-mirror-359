import os
import re
import logging
from typing import List, Dict, Any, Optional
import PyPDF2
import io
from datetime import datetime
import tiktoken

logger = logging.getLogger(__name__)

class TextProcessor:
    """Service for processing documents and creating text chunks for embedding."""
    
    def __init__(self):
        self.chunk_size = 800  # Hardcoded chunk size in tokens
        self.chunk_overlap = 100  # Hardcoded chunk overlap in tokens
        
        # Initialize tokenizer for accurate token counting
        try:
            self.tokenizer = tiktoken.get_encoding("cl100k_base")  # GPT-4 tokenizer
        except Exception as e:
            logger.warning(f"Failed to load tokenizer: {e}")
            self.tokenizer = None
    
    def extract_text_from_pdf(self, pdf_content: bytes) -> str:
        """Extract text content from PDF bytes."""
        try:
            pdf_file = io.BytesIO(pdf_content)
            pdf_reader = PyPDF2.PdfReader(pdf_file)
            
            text_content = ""
            for page_num, page in enumerate(pdf_reader.pages):
                try:
                    page_text = page.extract_text()
                    if page_text.strip():
                        text_content += f"\n--- Page {page_num + 1} ---\n"
                        text_content += page_text
                        text_content += "\n"
                except Exception as e:
                    logger.warning(f"Failed to extract text from page {page_num + 1}: {e}")
                    continue
            
            if not text_content.strip():
                raise ValueError("No text content extracted from PDF")
            
            logger.info(f"Extracted {len(text_content)} characters from PDF with {len(pdf_reader.pages)} pages")
            return text_content.strip()
            
        except Exception as e:
            logger.error(f"Failed to extract text from PDF: {e}")
            raise
    
    def extract_text_from_file(self, file_content: bytes, file_type: str) -> str:
        """Extract text from various file types."""
        try:
            if file_type == "application/pdf":
                return self.extract_text_from_pdf(file_content)
            elif file_type == "text/plain":
                return file_content.decode('utf-8', errors='ignore')
            elif file_type == "text/markdown":
                return file_content.decode('utf-8', errors='ignore')
            elif file_type == "application/json":
                return self._extract_from_json(file_content)
            elif file_type in ["text/csv", "application/csv"]:
                return self._extract_from_csv(file_content)
            elif file_type in ["application/vnd.openxmlformats-officedocument.spreadsheetml.sheet", "application/vnd.ms-excel"]:
                return self._extract_from_excel(file_content)
            elif file_type in ["application/msword", "application/vnd.openxmlformats-officedocument.wordprocessingml.document"]:
                return self._extract_from_word(file_content)
            else:
                raise ValueError(f"Unsupported file type: {file_type}")
                
        except Exception as e:
            logger.error(f"Failed to extract text from {file_type}: {e}")
            raise
    
    def preprocess_text(self, text: str) -> str:
        """Clean and preprocess text content."""
        if not text:
            return ""
        
        # Remove excessive whitespace
        text = re.sub(r'\s+', ' ', text)
        
        # Remove special characters but keep basic punctuation
        text = re.sub(r'[^\w\s\.\,\!\?\;\:\-\(\)\[\]\"\']+', ' ', text)
        
        # Remove very short lines that are likely artifacts
        lines = text.split('\n')
        cleaned_lines = []
        for line in lines:
            line = line.strip()
            if len(line) > 10:  # Keep lines with substantial content
                cleaned_lines.append(line)
        
        cleaned_text = ' '.join(cleaned_lines)
        
        # Final cleanup
        cleaned_text = re.sub(r'\s+', ' ', cleaned_text).strip()
        
        logger.debug(f"Preprocessed text: {len(text)} -> {len(cleaned_text)} characters")
        return cleaned_text
    
    def count_tokens(self, text: str) -> int:
        """Count tokens in text using tiktoken tokenizer."""
        if not self.tokenizer or not text:
            # Fallback estimation: ~4 characters per token
            return len(text) // 4
        
        try:
            return len(self.tokenizer.encode(text))
        except Exception as e:
            logger.warning(f"Token counting failed: {e}")
            return len(text) // 4
    
    def create_chunks(self, text: str, attachment_id: str, filename: str, 
                     conversation_id: str, file_type: str, space_id: str = None, 
                     user_id: str = None, project_id: str = None, scope: str = 'project') -> List[Dict[str, Any]]:
        """Split text into overlapping chunks for embedding."""
        if not text:
            return []
        
        # Preprocess text
        cleaned_text = self.preprocess_text(text)
        
        chunks = []
        chunk_index = 0
        start_pos = 0
        
        while start_pos < len(cleaned_text):
            # Extract chunk
            chunk_end = self._find_chunk_end(cleaned_text, start_pos)
            chunk_text = cleaned_text[start_pos:chunk_end].strip()
            
            if chunk_text and len(chunk_text) > 50:  # Only keep substantial chunks
                chunk_data = {
                    "conversation_id": conversation_id,
                    "attachment_id": attachment_id,
                    "space_id": space_id,  # Include space_id
                    "user_id": user_id,  # Include user_id
                    "project_id": project_id,  # Include project_id
                    "scope": scope,  # Include scope
                    "filename": filename,
                    "chunk_index": chunk_index,
                    "chunk_text": chunk_text,
                    "file_type": file_type,
                    "content_type": "text",  # Mark as text content
                    "created_at": datetime.utcnow().isoformat(),
                    "token_count": self.count_tokens(chunk_text)
                }
                chunks.append(chunk_data)
                chunk_index += 1
            
            # Move to next chunk with overlap
            next_start = self._calculate_next_start(cleaned_text, start_pos, chunk_end)
            if next_start <= start_pos:  # Avoid infinite loop
                break
            start_pos = next_start
        
        logger.info(f"Created {len(chunks)} chunks from {filename}")
        return chunks
    
    def _find_chunk_end(self, text: str, start_pos: int) -> int:
        """Find the end position for a chunk, respecting token limits."""
        # Start with character-based estimate
        estimated_end = start_pos + (self.chunk_size * 4)  # ~4 chars per token
        
        if estimated_end >= len(text):
            return len(text)
        
        # Try to break at sentence boundaries
        sentence_end = self._find_sentence_boundary(text, start_pos, estimated_end)
        if sentence_end > start_pos:
            chunk_text = text[start_pos:sentence_end]
            token_count = self.count_tokens(chunk_text)
            
            # Adjust if token count is too high or low
            if token_count <= self.chunk_size:
                return sentence_end
        
        # Fallback: use token-based chunking
        return self._find_token_boundary(text, start_pos)
    
    def _find_sentence_boundary(self, text: str, start_pos: int, estimated_end: int) -> int:
        """Find the nearest sentence boundary within the estimated range."""
        # Look for sentence endings near the estimated end
        search_range = 200  # Characters to search around estimated end
        search_start = max(start_pos + 100, estimated_end - search_range)
        search_end = min(len(text), estimated_end + search_range)
        
        search_text = text[search_start:search_end]
        
        # Find sentence endings (., !, ?)
        sentence_endings = []
        for match in re.finditer(r'[\.!?]\s+', search_text):
            pos = search_start + match.end()
            if pos > start_pos + 100:  # Ensure minimum chunk size
                sentence_endings.append(pos)
        
        if sentence_endings:
            # Return the ending closest to our estimated position
            return min(sentence_endings, key=lambda x: abs(x - estimated_end))
        
        return estimated_end
    
    def _find_token_boundary(self, text: str, start_pos: int) -> int:
        """Find chunk end based on exact token counting."""
        current_pos = start_pos
        max_pos = len(text)
        
        # Binary search for optimal position
        while current_pos < max_pos:
            test_end = min(current_pos + (self.chunk_size * 4), max_pos)
            chunk_text = text[start_pos:test_end]
            token_count = self.count_tokens(chunk_text)
            
            if token_count <= self.chunk_size:
                current_pos = test_end
                if test_end == max_pos:
                    break
            else:
                # Too many tokens, reduce by 10%
                test_end = start_pos + int((test_end - start_pos) * 0.9)
                return max(test_end, start_pos + 100)
        
        return current_pos
    
    def _calculate_next_start(self, text: str, current_start: int, current_end: int) -> int:
        """Calculate the start position for the next chunk with overlap."""
        if current_end >= len(text):
            return len(text)
        
        # Calculate overlap in characters (approximate)
        overlap_chars = self.chunk_overlap * 4  # ~4 chars per token
        next_start = max(current_start + 1, current_end - overlap_chars)
        
        # Try to start at a word boundary
        word_boundary = self._find_word_boundary(text, next_start)
        return word_boundary if word_boundary > current_start else next_start
    
    def _find_word_boundary(self, text: str, pos: int) -> int:
        """Find the nearest word boundary at or after the given position."""
        if pos >= len(text):
            return len(text)
        
        # Look for space or punctuation
        for i in range(pos, min(pos + 50, len(text))):
            if text[i] in ' \t\n.,!?;:':
                return i + 1
        
        return pos
    
    def get_processing_stats(self, chunks: List[Dict[str, Any]]) -> Dict[str, Any]:
        """Get statistics about the processed chunks."""
        if not chunks:
            return {"total_chunks": 0}
        
        total_tokens = sum(chunk.get("token_count", 0) for chunk in chunks)
        total_chars = sum(len(chunk.get("chunk_text", "")) for chunk in chunks)
        
        return {
            "total_chunks": len(chunks),
            "total_tokens": total_tokens,
            "total_characters": total_chars,
            "avg_tokens_per_chunk": total_tokens // len(chunks) if chunks else 0,
            "avg_chars_per_chunk": total_chars // len(chunks) if chunks else 0,
            "filename": chunks[0].get("filename", "unknown"),
            "file_type": chunks[0].get("file_type", "unknown")
        }

    def _extract_from_json(self, file_content: bytes) -> str:
        """Extract searchable text from JSON files."""
        try:
            import json
            
            json_content = file_content.decode('utf-8', errors='ignore')
            data = json.loads(json_content)
            
            # Convert JSON to searchable text
            searchable_text = self._json_to_text(data)
            
            logger.info(f"Extracted {len(searchable_text)} characters from JSON")
            return searchable_text
            
        except json.JSONDecodeError as e:
            logger.error(f"Failed to parse JSON: {e}")
            raise ValueError(f"Invalid JSON format: {e}")
        except Exception as e:
            logger.error(f"Failed to extract text from JSON: {e}")
            raise

    def _extract_from_csv(self, file_content: bytes) -> str:
        """Extract searchable text from CSV files."""
        try:
            import csv
            import io
            
            csv_content = file_content.decode('utf-8', errors='ignore')
            csv_reader = csv.reader(io.StringIO(csv_content))
            
            rows = []
            headers = None
            row_count = 0
            
            for row in csv_reader:
                if row_count == 0:
                    headers = row
                    rows.append(f"Headers: {', '.join(row)}")
                else:
                    # Create searchable row text with headers
                    if headers and len(row) >= len(headers):
                        row_text = []
                        for i, value in enumerate(row[:len(headers)]):
                            if value.strip():  # Only include non-empty values
                                row_text.append(f"{headers[i]}: {value}")
                        if row_text:
                            rows.append(f"Row {row_count}: {', '.join(row_text)}")
                    else:
                        # Fallback: just join row values
                        row_values = [val for val in row if val.strip()]
                        if row_values:
                            rows.append(f"Row {row_count}: {', '.join(row_values)}")
                
                row_count += 1
                # No row limit for indexing - process all data
            
            searchable_text = "\n".join(rows)
            logger.info(f"Extracted {len(searchable_text)} characters from CSV with {row_count} rows")
            return searchable_text
            
        except Exception as e:
            logger.error(f"Failed to extract text from CSV: {e}")
            raise

    def _extract_from_excel(self, file_content: bytes) -> str:
        """Extract searchable text from Excel files."""
        try:
            # Note: This requires openpyxl to be added to dependencies
            try:
                from openpyxl import load_workbook
            except ImportError:
                raise ValueError("openpyxl library required for Excel file processing. Add 'openpyxl>=3.1.0' to dependencies.")
            
            import io
            
            workbook = load_workbook(io.BytesIO(file_content), read_only=True)
            
            sheets_text = []
            
            for sheet_name in workbook.sheetnames:  # Process ALL sheets
                sheet = workbook[sheet_name]
                sheets_text.append(f"=== Sheet: {sheet_name} ===")
                
                rows_processed = 0
                headers = None
                
                for row in sheet.iter_rows(values_only=True):  # Process ALL rows
                    if rows_processed == 0:
                        # First row as headers
                        headers = [str(cell) if cell is not None else "" for cell in row]
                        sheets_text.append(f"Headers: {', '.join([h for h in headers if h])}")
                    else:
                        # Data rows
                        row_data = []
                        for i, cell in enumerate(row):
                            if cell is not None and str(cell).strip():
                                if headers and i < len(headers) and headers[i]:
                                    row_data.append(f"{headers[i]}: {cell}")
                                else:
                                    row_data.append(str(cell))
                        
                        if row_data:
                            sheets_text.append(f"Row {rows_processed}: {', '.join(row_data)}")
                    
                    rows_processed += 1
            
            searchable_text = "\n".join(sheets_text)
            logger.info(f"Extracted {len(searchable_text)} characters from Excel with {len(workbook.sheetnames)} sheets")
            return searchable_text
            
        except Exception as e:
            logger.error(f"Failed to extract text from Excel: {e}")
            raise

    def _extract_from_word(self, file_content: bytes) -> str:
        """Extract text from Word documents (DOC/DOCX)."""
        try:
            # Note: This requires python-docx to be added to dependencies
            try:
                from docx import Document
            except ImportError:
                raise ValueError("python-docx library required for Word document processing. Add 'python-docx>=1.1.0' to dependencies.")
            
            import io
            
            # Load document from bytes
            doc = Document(io.BytesIO(file_content))
            
            text_content = []
            
            # Extract text from paragraphs
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_content.append(paragraph.text.strip())
            
            # Extract text from tables
            for table in doc.tables:
                table_text = []
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        table_text.append(" | ".join(row_text))
                
                if table_text:
                    text_content.append("TABLE:")
                    text_content.extend(table_text)
                    text_content.append("END_TABLE")
            
            full_text = "\n".join(text_content)
            
            if not full_text.strip():
                raise ValueError("No text content extracted from Word document")
            
            logger.info(f"Extracted {len(full_text)} characters from Word document")
            return full_text
            
        except Exception as e:
            logger.error(f"Failed to extract text from Word document: {e}")
            raise

    def _json_to_text(self, data, prefix="") -> str:
        """Convert JSON data to searchable text."""
        text_parts = []
        
        if isinstance(data, dict):
            for key, value in data.items():
                current_prefix = f"{prefix}.{key}" if prefix else key
                
                if isinstance(value, (dict, list)):
                    text_parts.append(f"{current_prefix}:")
                    text_parts.append(self._json_to_text(value, current_prefix))
                else:
                    text_parts.append(f"{current_prefix}: {value}")
                    
        elif isinstance(data, list):
            for i, item in enumerate(data):
                current_prefix = f"{prefix}[{i}]" if prefix else f"item_{i}"
                text_parts.append(self._json_to_text(item, current_prefix))
                
        else:
            return str(data)
        
        return "\n".join(text_parts)

 