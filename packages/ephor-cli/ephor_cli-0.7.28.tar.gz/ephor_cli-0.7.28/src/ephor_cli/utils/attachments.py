import base64
import logging
from typing import List, Dict, Any

from ephor_cli.services.conversation_attachment import ConversationAttachmentService
from langchain_core.messages import AIMessage

logger = logging.getLogger(__name__)

def process_attachments(message, s3_service, user_id: str = None, project_id: str = None, conversation_id: str = None, space_id: str = None):
    """Process both message-level and conversation-level attachments."""
    attachments = message.additional_kwargs.get("attachments", [])
    if not attachments and not (user_id and project_id and conversation_id):
        return message

    # Handle both string and list content formats
    if isinstance(message.content, str):
        content = [{"type": "text", "text": message.content}]
    elif isinstance(message.content, list):
        content = message.content.copy()  
    else:
        content = [{"type": "text", "text": ""}]

    # Process message-level attachments
    content = _process_message_level_attachments(attachments, content, s3_service)

    # Process conversation-level attachments with intelligent indexing
    if user_id and project_id and conversation_id:
        content = _process_conversation_level_attachments(
            content, message, s3_service, user_id, project_id, conversation_id, space_id
        )

    message.content = content
    logger.debug("[attachments.py] Message content after intelligent attachment processing: %s", len(content))
    return message

def _process_message_level_attachments(attachments: List[Dict], content: List[Dict], s3_service) -> List[Dict]:
    """Process message-level attachments with direct content injection."""
    for att in attachments:
        s3_key = att.get("s3_key")
        file_type = att.get("type", "")
        name = att.get("name", "")
        if not s3_key:
            continue
        try:
            file_content = s3_service.get_file_content(s3_key, use_cache=False)  # No cache for message-level (attachments are only sent with message therfore no caching)
            if not file_content:
                continue

            if file_type.startswith("image/"):
                b64_data = base64.b64encode(file_content).decode("utf-8")
                content.append({
                    "type": "image",
                    "source_type": "base64",
                    "data": b64_data,
                    "mime_type": file_type,
                    "filename": name
                })
            elif file_type == "application/pdf":
                b64_data = base64.b64encode(file_content).decode("utf-8")
                content.append({
                    "type": "file",
                    "source_type": "base64",
                    "data": b64_data,
                    "mime_type": "application/pdf",
                    "filename": name
                })
            elif file_type == "text/plain":
                try:
                    text_content = file_content.decode("utf-8")
                    content.append({
                        "type": "text",
                        "text": f"<context>Content of {name}:\n\n{text_content}</context>"
                    })
                except UnicodeDecodeError:
                    pass
            elif file_type == "text/markdown":
                try:
                    text_content = file_content.decode("utf-8")
                    content.append({
                        "type": "text",
                        "text": f"<context>Markdown content of {name}:\n\n{text_content}</context>"
                    })
                except UnicodeDecodeError:
                    pass
            elif file_type == "application/json":
                try:
                    import json
                    json_data = json.loads(file_content.decode('utf-8'))
                    formatted_json = json.dumps(json_data, indent=2)
                    content.append({
                        "type": "text",
                        "text": f"<context>JSON content of {name}:\n\n{formatted_json}</context>"
                    })
                except (json.JSONDecodeError, UnicodeDecodeError):
                    content.append({
                        "type": "text", 
                        "text": f"[JSON file {name} could not be parsed]"
                    })
            elif file_type in ["text/csv", "application/csv"]:
                try:
                    import csv
                    import io
                    csv_content = file_content.decode('utf-8')
                    csv_reader = csv.reader(io.StringIO(csv_content))
                    rows = list(csv_reader)[:50]  # Limit to first 50 rows for direct display
                    formatted_csv = "\n".join(["\t".join(row) for row in rows])
                    content.append({
                        "type": "text",
                        "text": f"<context>CSV content of {name} (first 50 rows):\n\n{formatted_csv}</context>"
                    })
                except UnicodeDecodeError:
                    content.append({
                        "type": "text",
                        "text": f"[CSV file {name} could not be parsed]"
                    })
            elif file_type in ["application/msword", "application/vnd.openxmlformats-officedocument.wordprocessingml.document"]:
                try:
                    from docx import Document
                    import io
                    
                    doc = Document(io.BytesIO(file_content))
                    word_content = []
                    
                    # Extract paragraphs
                    for paragraph in doc.paragraphs:
                        if paragraph.text.strip():
                            word_content.append(paragraph.text.strip())
                    
                    # Extract tables
                    for table in doc.tables:
                        word_content.append("TABLE:")
                        for row in table.rows:
                            row_text = []
                            for cell in row.cells:
                                if cell.text.strip():
                                    row_text.append(cell.text.strip())
                            if row_text:
                                word_content.append(" | ".join(row_text))
                        word_content.append("END_TABLE")
                    
                    if word_content:
                        content.append({
                            "type": "text",
                            "text": f"<context>Word document content of {name}:\n\n{chr(10).join(word_content)}</context>"
                        })
                    else:
                        content.append({
                            "type": "text",
                            "text": f"[Word document {name} contains no extractable text]"
                        })
                except Exception as e:
                    content.append({
                        "type": "text",
                        "text": f"[Word document {name} could not be processed: {e}]"
                    })
            elif file_type in ["application/vnd.openxmlformats-officedocument.spreadsheetml.sheet", "application/vnd.ms-excel"]:
                try:
                    from openpyxl import load_workbook
                    import io
                    
                    workbook = load_workbook(io.BytesIO(file_content), read_only=True)
                    excel_content = []
                    
                    for sheet_name in workbook.sheetnames[:2]:  # Limit to first 2 sheets for direct display
                        sheet = workbook[sheet_name]
                        excel_content.append(f"=== Sheet: {sheet_name} ===")
                        
                        row_count = 0
                        for row in sheet.iter_rows(values_only=True, max_row=50):  # Limit to 50 rows
                            if row_count == 0:
                                headers = [str(cell) if cell is not None else "" for cell in row]
                                excel_content.append("Headers: " + "\t".join(headers))
                            else:
                                row_data = [str(cell) if cell is not None else "" for cell in row]
                                excel_content.append(f"Row {row_count}: " + "\t".join(row_data))
                            row_count += 1
                        
                        if row_count >= 50:
                            excel_content.append("... (truncated for display)")
                    
                    content.append({
                        "type": "text",
                        "text": f"<context>Excel content of {name}:\n\n{chr(10).join(excel_content)}</context>"
                    })
                except Exception as e:
                    content.append({
                        "type": "text",
                        "text": f"[Excel file {name} could not be processed: {e}]"
                    })
            else:
                content.append({
                    "type": "text",
                    "text": f"[File {name} attached, but type {file_type} is not supported for LLM input]"
                })
        except Exception as e:
            content.append({
                "type": "text",
                "text": f"[Could not load {name}: {e}]"
            })
    
    return content

def _process_conversation_level_attachments(content: List[Dict], message, s3_service, 
                                          user_id: str, project_id: str, conversation_id: str, space_id: str = None) -> List[Dict]:
    """Process conversation-level attachments with intelligent indexing strategy."""
    try:
        # Try to use indexing system for semantic search
        relevant_chunks = _get_relevant_chunks_via_indexing(message, conversation_id, user_id, project_id, space_id)
        
        if relevant_chunks:
            logger.info(f"[attachments.py] Using indexing: Retrieved {len(relevant_chunks)} relevant chunks")
            
            # Track which images were found relevant via their captions
            relevant_image_attachments = set()
            
            # Add retrieved chunks as context
            for chunk in relevant_chunks:
                content.append({
                    "type": "text", 
                    "text": f"<retrieved_context source=\"{chunk.get('filename', 'unknown')}\">\n{chunk.get('chunk_text', '')}\n</retrieved_context>"
                })
                
                # Track if this chunk is an image caption
                if chunk.get('content_type') == 'image_caption':
                    relevant_image_attachments.add(chunk.get('attachment_id'))
            
            # Only add raw images if their captions were found semantically relevant
            if relevant_image_attachments:
                added_count = _add_scope_aware_images_only(content, s3_service, relevant_chunks, relevant_image_attachments)
                logger.info(f"[attachments.py] Added {added_count} semantically relevant images from scope-aware search")
            
            return content
        else:
            logger.info("[attachments.py] No relevant chunks found via indexing - respecting semantic filtering, sending no attachment context")
            # Do NOT fall back to sending all attachments - respect the semantic filtering decision
            # If no chunks are semantically relevant, then no attachments should be sent to the agent
            return content
            
    except Exception as e:
        logger.warning(f"[attachments.py] Indexing failed: {e}, falling back to direct processing")
    
    # Fallback: Use traditional direct content injection ONLY when indexing system fails, not when it works but finds nothing
    logger.info("[attachments.py] Indexing system failed, using direct attachment processing as fallback")
    return _process_conversation_attachments_directly(content, s3_service, user_id, project_id, conversation_id)

def _get_relevant_chunks_via_indexing(message, conversation_id: str, user_id: str, project_id: str, space_id: str = None) -> List[Dict[str, Any]]:
    """Get relevant chunks using the indexing system."""
    try:
        from ephor_cli.services.conversation_indexing_service import ConversationIndexingService
        
        # Form query context: current message + previous response (if available)
        query_text = _extract_query_text(message)
        previous_response = _extract_previous_response(message, conversation_id, user_id, project_id)
        
        # Fallback for backward compatibility: check message if space_id not provided as parameter
        if not space_id:
            space_id = message.additional_kwargs.get("space_id")
        
        search_scope = f"space_id: {space_id}" if space_id else f"conversation_id: {conversation_id}"
        logger.info(f"[attachments.py] Searching with query: '{query_text}' in {search_scope}")
        
        # Use indexing service for semantic search with space_id priority
        indexing_service = ConversationIndexingService()
        relevant_chunks = indexing_service.search_conversation_content(
            conversation_id=conversation_id,
            query_text=query_text,
            previous_response=previous_response,
            space_id=space_id,
            user_id=user_id,
            project_id=project_id
        )
        
        return relevant_chunks
        
    except ImportError:
        logger.warning("[attachments.py] Indexing service not available")
        return []
    except Exception as e:
        logger.error(f"[attachments.py] Error in semantic search: {e}")
        return []

def _extract_query_text(message) -> str:
    """Extract text from current message for query formation."""
    if isinstance(message.content, str):
        return message.content
    elif isinstance(message.content, list):
        text_parts = []
        for item in message.content:
            if item.get("type") == "text":
                text_parts.append(item.get("text", ""))
        return " ".join(text_parts)
    return ""

def _extract_previous_response(message, conversation_id: str, user_id: str, project_id: str) -> str:
    """Extract previous agent response for query context."""
    try:
        # Get the last message ID from the current message's metadata
        last_message_id = message.additional_kwargs.get("last_message_id")
        if not last_message_id:
            logger.info(f"[attachments.py] No last message ID found for message: {message.id}")
            return ""
            
        # Get the previous message directly using the last_message_id and conversation_id
        from ephor_cli.services import message_service
        previous_message = message_service.get_message(
            user_id,
            project_id,
            conversation_id,
            last_message_id
        )
        
        if not previous_message:
            logger.warning(f"[attachments.py] No previous message found with ID: {last_message_id}")
            return ""
            
        # Only use content from AI messages
        if not isinstance(previous_message, AIMessage):
            logger.info(f"[attachments.py] Previous message is not from AI, skipping")
            return ""
            
        # Extract text content from the message
        content = previous_message.content
        if isinstance(content, str):
            return content
        elif isinstance(content, list):
            # Handle list format content
            text_parts = []
            for item in content:
                if item.get("type") == "text":
                    text_parts.append(item.get("text", ""))
            return " ".join(text_parts)
        return ""
        
    except Exception as e:
        logger.error(f"[attachments.py] Error extracting previous response: {e}")
        return ""

def _add_conversation_images_directly(content: List[Dict], s3_service, user_id: str, project_id: str, conversation_id: str):
    """Add ALL conversation-level images directly (fallback method when indexing fails)."""
    try:
        conversation_attachment_service = ConversationAttachmentService()
        conversation_attachments = conversation_attachment_service.list_attachments(
            user_id, project_id, conversation_id
        )

        for att in conversation_attachments:
            # Only process images - PDFs/text are handled by indexing (fallback mode)
            if att.file_type.startswith("image/"):
                try:
                    file_content = s3_service.get_file_content(att.s3_key, use_cache=True)  # Use cache for conversation-level
                    if file_content:
                        b64_data = base64.b64encode(file_content).decode("utf-8")
                        content.append({
                            "type": "image",
                            "source_type": "base64",
                            "data": b64_data,
                            "mime_type": att.file_type,
                            "filename": att.file_name
                        })
                        logger.debug(f"[attachments.py] Added conversation image: {att.file_name}")
                except Exception as e:
                    logger.error(f"[attachments.py] Error processing conversation image {att.file_name}: {e}")
                    
    except Exception as e:
        logger.error(f"[attachments.py] Error accessing conversation images: {e}")

def _add_scope_aware_images_only(content: List[Dict], s3_service, relevant_chunks: List[Dict], relevant_attachment_ids: set) -> int:
    """Add only image files for semantically relevant attachments from scope-aware search.
    
    This follows the main branch approach: only pass chunk text content to agent,
    but add actual image files if their captions were found semantically relevant.
    
    Args:
        content: List to append image content to
        s3_service: S3 service for downloading images
        relevant_chunks: List of relevant chunks with metadata
        relevant_attachment_ids: Set of attachment IDs that were found semantically relevant
        
    Returns:
        Number of images successfully added
    """
    from ephor_cli.services.conversation_attachment import ConversationAttachmentService
    
    try:
        conversation_attachment_service = ConversationAttachmentService()
        
        # Create attachment metadata map
        attachment_metadata = {}
        for chunk in relevant_chunks:
            att_id = chunk.get('attachment_id')
            if att_id and att_id in relevant_attachment_ids:
                attachment_metadata[att_id] = chunk
        
        images_added = 0
        
        # Process each relevant attachment, but only add images
        for att_id in relevant_attachment_ids:
            chunk_meta = attachment_metadata.get(att_id, {})
            
            try:
                # Get attachment details from the original conversation
                original_conversation_id = chunk_meta.get('conversation_id')
                original_user_id = chunk_meta.get('user_id')  
                original_project_id = chunk_meta.get('project_id')
                
                if not all([original_conversation_id, original_user_id, original_project_id]):
                    logger.warning(f"[attachments.py] Missing metadata for attachment {att_id}, skipping")
                    continue
                
                # Get the attachment from its original conversation
                attachments = conversation_attachment_service.list_attachments(
                    original_user_id, original_project_id, original_conversation_id
                )
                attachment = next((att for att in attachments if att.id == att_id), None)
                
                if not attachment:
                    logger.warning(f"[attachments.py] Attachment {att_id} not found in conversation {original_conversation_id}")
                    continue
                
                # Only process images (following main branch approach)
                if attachment.file_type.startswith("image/"):
                    file_content = s3_service.get_file_content(attachment.s3_key, use_cache=True)
                    if not file_content:
                        logger.warning(f"[attachments.py] Could not download image for attachment {att_id}")
                        continue
                    
                    # Add the image
                    b64_data = base64.b64encode(file_content).decode("utf-8")
                    content.append({
                        "type": "image",
                        "source_type": "base64",
                        "data": b64_data,
                        "mime_type": attachment.file_type,
                        "filename": attachment.file_name
                    })
                    
                    images_added += 1
                    logger.debug(f"[attachments.py] Added scope-aware image: {attachment.file_name} (from conversation {original_conversation_id})")
                
                # For non-images: chunk text content is already added as <retrieved_context>,                         
            except Exception as e:
                logger.error(f"[attachments.py] Error processing scope-aware image {att_id}: {e}")
                    
        logger.info(f"[attachments.py] Successfully added {images_added} scope-aware images")
        return images_added
                    
    except Exception as e:
        logger.error(f"[attachments.py] Error in scope-aware image processing: {e}")
        return 0

def _add_specific_conversation_images(content: List[Dict], s3_service, user_id: str, project_id: str, 
                                    conversation_id: str, relevant_attachment_ids: set):
    """Add only specific semantically relevant conversation images.
    
    Args:
        content: List to append image content to
        s3_service: S3 service for downloading images
        user_id: User ID
        project_id: Project ID
        conversation_id: Conversation ID
        relevant_attachment_ids: Set of attachment IDs that were found semantically relevant
    """
    try:
        conversation_attachment_service = ConversationAttachmentService()
        conversation_attachments = conversation_attachment_service.list_attachments(
            user_id, project_id, conversation_id
        )

        for att in conversation_attachments:
            # Only process images that were found semantically relevant
            if att.file_type.startswith("image/") and att.id in relevant_attachment_ids:
                try:
                    file_content = s3_service.get_file_content(att.s3_key,use_cache=True)
                    if file_content:
                        b64_data = base64.b64encode(file_content).decode("utf-8")
                        content.append({
                            "type": "image",
                            "source_type": "base64",
                            "data": b64_data,
                            "mime_type": att.file_type,
                            "filename": att.file_name
                        })
                        logger.debug(f"[attachments.py] Added semantically relevant image: {att.file_name}")
                except Exception as e:
                    logger.error(f"[attachments.py] Error processing relevant image {att.file_name}: {e}")
                    
    except Exception as e:
        logger.error(f"[attachments.py] Error accessing relevant conversation images: {e}")

def _process_conversation_attachments_directly(content: List[Dict], s3_service, user_id: str, project_id: str, conversation_id: str) -> List[Dict]:
    """Fallback: Process conversation attachments with direct content injection (original behavior)."""
    conversation_attachment_service = ConversationAttachmentService()
    conversation_attachments = conversation_attachment_service.list_attachments(
        user_id, project_id, conversation_id
    )

    for att in conversation_attachments:
        try:
            file_content = s3_service.get_file_content(att.s3_key, use_cache=True)
            if not file_content:
                continue

            if att.file_type.startswith("image/"):
                b64_data = base64.b64encode(file_content).decode("utf-8")
                content.append({
                    "type": "image",
                    "source_type": "base64",
                    "data": b64_data,
                    "mime_type": att.file_type,
                    "filename": att.file_name
                })
            elif att.file_type == "application/pdf":
                b64_data = base64.b64encode(file_content).decode("utf-8")
                content.append({
                    "type": "file",
                    "source_type": "base64",
                    "data": b64_data,
                    "mime_type": "application/pdf",
                    "filename": att.file_name
                })
            elif att.file_type == "text/plain":
                try:
                    text_content = file_content.decode("utf-8")
                    content.append({
                        "type": "text",
                        "text": f"<context>Content of {att.file_name}:\n\n{text_content}</context>"
                    })
                except UnicodeDecodeError:
                    pass
            elif att.file_type == "text/markdown":
                try:
                    text_content = file_content.decode("utf-8")
                    content.append({
                        "type": "text",
                        "text": f"<context>Markdown content of {att.file_name}:\n\n{text_content}</context>"
                    })
                except UnicodeDecodeError:
                    pass
            elif att.file_type == "application/json":
                try:
                    import json
                    json_data = json.loads(file_content.decode('utf-8'))
                    formatted_json = json.dumps(json_data, indent=2)
                    content.append({
                        "type": "text",
                        "text": f"<context>JSON content of {att.file_name}:\n\n{formatted_json}</context>"
                    })
                except (json.JSONDecodeError, UnicodeDecodeError):
                    content.append({
                        "type": "text", 
                        "text": f"[JSON file {att.file_name} could not be parsed]"
                    })
            elif att.file_type in ["text/csv", "application/csv"]:
                try:
                    import csv
                    import io
                    csv_content = file_content.decode('utf-8')
                    csv_reader = csv.reader(io.StringIO(csv_content))
                    rows = list(csv_reader)[:50]  # Limit to first 50 rows for direct display
                    formatted_csv = "\n".join(["\t".join(row) for row in rows])
                    content.append({
                        "type": "text",
                        "text": f"<context>CSV content of {att.file_name} (first 50 rows):\n\n{formatted_csv}</context>"
                    })
                except UnicodeDecodeError:
                    content.append({
                        "type": "text",
                        "text": f"[CSV file {att.file_name} could not be parsed]"
                    })
            elif att.file_type in ["application/msword", "application/vnd.openxmlformats-officedocument.wordprocessingml.document"]:
                try:
                    from docx import Document
                    import io
                    
                    doc = Document(io.BytesIO(file_content))
                    word_content = []
                    
                    # Extract paragraphs
                    for paragraph in doc.paragraphs:
                        if paragraph.text.strip():
                            word_content.append(paragraph.text.strip())
                    
                    # Extract tables
                    for table in doc.tables:
                        word_content.append("TABLE:")
                        for row in table.rows:
                            row_text = []
                            for cell in row.cells:
                                if cell.text.strip():
                                    row_text.append(cell.text.strip())
                            if row_text:
                                word_content.append(" | ".join(row_text))
                        word_content.append("END_TABLE")
                    
                    if word_content:
                        content.append({
                            "type": "text",
                            "text": f"<context>Word document content of {att.file_name}:\n\n{chr(10).join(word_content)}</context>"
                        })
                    else:
                        content.append({
                            "type": "text",
                            "text": f"[Word document {att.file_name} contains no extractable text]"
                        })
                except Exception as e:
                    content.append({
                        "type": "text",
                        "text": f"[Word document {att.file_name} could not be processed: {e}]"
                    })
            elif att.file_type in ["application/vnd.openxmlformats-officedocument.spreadsheetml.sheet", "application/vnd.ms-excel"]:
                try:
                    from openpyxl import load_workbook
                    import io
                    
                    workbook = load_workbook(io.BytesIO(file_content), read_only=True)
                    excel_content = []
                    
                    for sheet_name in workbook.sheetnames[:2]:  # Limit to first 2 sheets for direct display
                        sheet = workbook[sheet_name]
                        excel_content.append(f"=== Sheet: {sheet_name} ===")
                        
                        row_count = 0
                        for row in sheet.iter_rows(values_only=True, max_row=50):  # Limit to 50 rows
                            if row_count == 0:
                                headers = [str(cell) if cell is not None else "" for cell in row]
                                excel_content.append("Headers: " + "\t".join(headers))
                            else:
                                row_data = [str(cell) if cell is not None else "" for cell in row]
                                excel_content.append(f"Row {row_count}: " + "\t".join(row_data))
                            row_count += 1
                        
                        if row_count >= 50:
                            excel_content.append("... (truncated for display)")
                    
                    content.append({
                        "type": "text",
                        "text": f"<context>Excel content of {att.file_name}:\n\n{chr(10).join(excel_content)}</context>"
                    })
                except Exception as e:
                    content.append({
                        "type": "text",
                        "text": f"[Excel file {att.file_name} could not be processed: {e}]"
                    })
            else:
                content.append({
                    "type": "text",
                    "text": f"[File {att.file_name} attached, but type {att.file_type} is not supported for LLM input]"
                })
        except Exception as e:
            content.append({
                "type": "text",
                "text": f"[Could not load {att.file_name}: {e}]"
            })

    return content

def cleanup_text_content_for_storage(message):
    """Clean up text file content from message before database storage.
    
    Removes content within <context></context> tags and replaces with metadata to prevent database bloat.
    """
    if not isinstance(message.content, list):
        return message

    cleaned_content = []

    for item in message.content:
        if item.get("type") == "text":
            text = item.get("text", "")

            # Check if this contains context tags with content
            if ("<context>" in text and "</context>" in text) or \
               ("<retrieved_context" in text and "</retrieved_context>" in text):
                # Skip context content to prevent database bloat
                continue

        # Keep all other content as-is
        cleaned_content.append(item)

    # Update message content with cleaned version
    message.content = cleaned_content
    return message 